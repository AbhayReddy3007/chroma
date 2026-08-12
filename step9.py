"""
step9.py — Invalidity Claim Chart Assembler (Final Deliverable)
================================================================
Renders the final claim charts in Gland Pharma Ltd. Appendix A format.
Combines step7 (skeleton), step8a (evidence), step8b (grounds) into
editable Word tables + underlying JSON.

No new evidence. No invalidity narrative. Chart creation only.

Data sources:
  - Step 7:  claim skeletons (verbatim claims + limitations)
  - Step 8a: per-limitation evidence (passages with loci)
  - Step 8b: grounds (102/103 sets), gap list, grace annex
  - ChromaDB: patent chunks (for supplementary context if needed)
  - Analysis cache: blocking analyser results (for metadata)

Model: gemini-2.5-flash-preview-05-20 (for chart assembly only —
       formatting/structuring, NOT new evidence)

Usage:
    python step9.py --drug Axitinib                       # all patents
    python step9.py --drug Axitinib --patent US10123456   # one patent
    python step9.py --drug Axitinib --rerun_all           # cascade re-run 7→8a→8b→9
"""

from __future__ import annotations

import argparse
import asyncio
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

from google import genai
from google.genai import types

# ─────────────────────────────────────────────────────────────
# Paths
# ─────────────────────────────────────────────────────────────

STEP7_OUTPUT_DIR   = Path(os.getenv("STEP7_OUTPUT_DIR",   Path(__file__).parent / "step7_output"))
STEP8A_OUTPUT_DIR  = Path(os.getenv("STEP8A_OUTPUT_DIR",  Path(__file__).parent / "step8a_output"))
STEP8B_OUTPUT_DIR  = Path(os.getenv("STEP8B_OUTPUT_DIR",  Path(__file__).parent / "step8b_output"))
STEP9_OUTPUT_DIR   = Path(os.getenv("STEP9_OUTPUT_DIR",   Path(__file__).parent / "step9_output"))
CHROMA_DB_PATH     = str(Path(__file__).parent / "chroma_patent_db")
ANALYSIS_CACHE_DIR = Path(os.getenv("ANALYSIS_CACHE_DIR", Path(__file__).parent / "analysis_cache"))

# ─────────────────────────────────────────────────────────────
# Gemini — lazy init so missing API key doesn't crash on import
# (step9 itself makes no Gemini calls; the import is for type
#  hints only — but step8b bootstrapping may need it)
# ─────────────────────────────────────────────────────────────

MODEL          = "gemini-2.5-flash-preview-05-20"
_gemini_client = None

def _get_gemini_client():
    global _gemini_client
    if _gemini_client is None:
        api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise ValueError(
                "GOOGLE_API_KEY (or GEMINI_API_KEY) is not set.\n"
                "Step 9 itself makes no Gemini calls, but upstream steps (7, 8a, 8b) do.\n"
                "Set it in your .env or environment before running."
            )
        _gemini_client = genai.Client(api_key=api_key)
    return _gemini_client

# ─────────────────────────────────────────────────────────────
# ChromaDB + analysis cache (reuse helpers)
# ─────────────────────────────────────────────────────────────

_chroma_client = None

def _get_chroma():
    global _chroma_client
    if _chroma_client is None:
        import chromadb
        _chroma_client = chromadb.PersistentClient(path=CHROMA_DB_PATH)
    return _chroma_client


def _sanitize_collection_name(drug_name: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._-]", "_", drug_name.strip())
    safe = re.sub(r"[_\-]{2,}", "_", safe)
    safe = re.sub(r"^[^a-zA-Z0-9]+", "", safe)
    safe = re.sub(r"[^a-zA-Z0-9]+$", "", safe)
    safe = safe.ljust(3, "x")[:55]
    safe = re.sub(r"[^a-zA-Z0-9]+$", "", safe)
    return f"patents_{safe}"


def _load_analysis_cache(drug_name: str, patent_number: str) -> Optional[dict]:
    safe_drug = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name.strip().lower())
    cache_dir = ANALYSIS_CACHE_DIR / safe_drug
    if not cache_dir.exists():
        return None
    pn_norm = patent_number.replace("-", "").replace(" ", "").upper()
    for f in cache_dir.glob("*.json"):
        fn_norm = f.stem.replace("-", "").replace(" ", "").replace("_", "").upper()
        if pn_norm in fn_norm or fn_norm in pn_norm:
            try:
                return json.loads(f.read_text(encoding="utf-8"))
            except Exception:
                continue
    return None


# ─────────────────────────────────────────────────────────────
# Upstream data loaders
# ─────────────────────────────────────────────────────────────

def _load_json(path: Path) -> Optional[dict | list]:
    if not path.exists():
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"[Step 9] Failed to read {path.name}: {e}")
        return None


def _load_all_inputs(drug_name: str, patent_number: str) -> dict:
    """Load step7 skeleton, step8a evidence, and step8b grounds for one patent."""
    safe_drug   = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name)
    safe_patent = re.sub(r"[^a-zA-Z0-9_-]", "_", patent_number)

    skeleton = _load_json(STEP7_OUTPUT_DIR  / f"{safe_drug}_{safe_patent}_claim_skeleton.json")
    evidence = _load_json(STEP8A_OUTPUT_DIR / f"{safe_drug}_{safe_patent}_prior_art.json")
    grounds  = _load_json(STEP8B_OUTPUT_DIR / f"{safe_drug}_{safe_patent}_grounds.json")
    analysis = _load_analysis_cache(drug_name, patent_number)

    return {
        "skeleton": skeleton,
        "evidence": evidence,
        "grounds":  grounds,
        "analysis_cache": analysis,
    }


def _ensure_step8b(drug_name: str, force: bool = False) -> None:
    safe     = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name)
    existing = list(STEP8B_OUTPUT_DIR.glob(f"{safe}_*_grounds.json"))
    existing = [f for f in existing if "_all_grounds" not in f.name]
    if existing and not force:
        return
    print(f"[Step 9] Step 8b output not found — running step8b for '{drug_name}'...")
    try:
        from step8b import run_for_drug
    except ImportError:
        import importlib.util
        spec = importlib.util.spec_from_file_location("step8b", Path(__file__).parent / "step8b.py")
        mod  = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)
        run_for_drug = mod.run_for_drug
    # step8b.run_for_drug is now async
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        loop = None
    if loop and loop.is_running():
        asyncio.ensure_future(run_for_drug(drug_name=drug_name))
    else:
        asyncio.run(run_for_drug(drug_name=drug_name))


# ─────────────────────────────────────────────────────────────
# Deterministic chart assembly (no LLM needed)
# ─────────────────────────────────────────────────────────────

def _assemble_charts(
    patent_number: str,
    priority_date: str,
    skeleton:      dict,
    evidence:      dict,
    grounds_data:  dict,
) -> dict:
    """
    Assemble the final claim charts from step 7/8a/8b data.
    Purely deterministic — no LLM call.

    Returns the full step9 JSON structure.
    """
    # Index evidence by (claim_number, limitation_id)
    ev_index: dict[tuple, dict] = {}
    for lr in evidence.get("limitation_results", []):
        key = (lr.get("claim_number"), lr.get("limitation_id"))
        ev_index[key] = lr

    # Collect all references for the reference list
    all_refs: dict[str, dict] = {}
    for lr in evidence.get("limitation_results", []):
        for ev in lr.get("evidence", []):
            ref_id = ev.get("reference_id", "")
            if ref_id and ref_id not in all_refs:
                all_refs[ref_id] = {
                    "short_name":       ref_id,
                    "full_citation":    f"{ref_id} ({ev.get('source', '?')})",
                    "publication_date": ev.get("publication_date", "?"),
                    "pre_priority":     ev.get("pre_priority", False),
                    "access":           "publicly accessible",
                }

    charts:           list = []
    coverage_summary: list = []
    all_gap_lims:     list = []
    all_grace_annex:  list = []

    for claim_ground in grounds_data.get("claim_grounds", []):
        cn    = claim_ground["claim_number"]
        total = claim_ground["total_limitations"]

        # Find the skeleton claim
        skel_claim = None
        for sc in skeleton.get("independent_claims", []):
            if sc.get("claim_number") == cn:
                skel_claim = sc
                break

        if not skel_claim:
            continue

        limitations = skel_claim.get("limitations", [])

        # Build a chart for each ground
        for ground in claim_ground.get("grounds", []):
            ground_id  = ground["ground_id"]
            basis      = ground["basis"]
            refs       = ground["references"]
            covered_set = set(ground.get("covered_lims", []))

            # Build basis line
            if basis == "102":
                basis_line = (
                    f'"{refs[0]}" anticipates Claim {cn} under 35 U.S.C. §102'
                )
            else:
                if len(refs) == 1:
                    basis_line = f'"{refs[0]}" renders Claim {cn} obvious under 35 U.S.C. §103'
                elif len(refs) == 2:
                    basis_line = (
                        f'"{refs[0]}" in view of "{refs[1]}" renders '
                        f"Claim {cn} obvious under 35 U.S.C. §103"
                    )
                else:
                    combo = f'"{refs[0]}" in view of ' + " and ".join(
                        f'"{r}"' for r in refs[1:]
                    )
                    basis_line = f"{combo} renders Claim {cn} obvious under 35 U.S.C. §103"

            # Build rows (one per limitation)
            rows = []
            for lim in limitations:
                lid   = lim["limitation_id"]
                ltext = lim["limitation_text_verbatim"]

                # Find evidence passages for this limitation from ground refs
                cells = []
                ev_data = ev_index.get((cn, lid))
                if ev_data and lid in covered_set:
                    for ref_id in refs:
                        passages = []
                        for ev in ev_data.get("evidence", []):
                            if ev.get("reference_id") == ref_id:
                                passage = ev.get("passage_verbatim", "")
                                locus   = ev.get("locus", "")
                                if passage:
                                    passages.append({
                                        "passage_verbatim": passage,
                                        "locus":            locus,
                                    })
                        if passages:
                            cells.append({
                                "reference_id": ref_id,
                                "passages":     passages,
                            })

                if not cells and lid not in covered_set:
                    cells.append({
                        "reference_id": "—",
                        "passages":     [{"passage_verbatim": "Not disclosed — see Gap List",
                                          "locus": ""}],
                    })

                rows.append({
                    "limitation_id":   lid,
                    "limitation_text": ltext,
                    "cells":           cells,
                })

            charts.append({
                "ground_id":              ground_id,
                "basis":                  basis,
                "references":             refs,
                "claim_number":           cn,
                "basis_line":             basis_line,
                "combination_rationale":  ground.get("combination_rationale", ""),
                "rows":                   rows,
            })

            # Coverage summary entry
            uncov_ids = [lim["limitation_id"] for lim in limitations if lim["limitation_id"] not in covered_set]
            coverage_summary.append({
                "claim":                cn,
                "ground":              f"§{basis}",
                "references":          refs,
                "covered_total":       f"{len(covered_set)}/{total}",
                "coverage_pct":        ground["coverage_pct"],
                "uncovered_lim_ids":   uncov_ids,
                "combination_rationale": ground.get("combination_rationale", ""),
            })

        # Strength ranking (from step8b LLM)
        strength_ranking = claim_ground.get("strength_ranking", [])

        # Gap list — now includes recommended_source from step8b LLM
        for gap in claim_ground.get("gap_limitations", []):
            all_gap_lims.append({
                "claim_number":       cn,
                "limitation_id":      gap["limitation_id"],
                "limitation_text":    gap.get("limitation_text", ""),
                "flags":              gap.get("flags", []),
                "recommended_source": gap.get("recommended_source", ""),
                "recommendation":     gap.get("recommendation", ""),
            })

        # Grace annex
        for ga in claim_ground.get("grace_only_limitations", []):
            all_grace_annex.append({
                "claim_number":   cn,
                "limitation_id":  ga["limitation_id"],
                "reference_id":   ga["reference_id"],
            })

    return {
        "patent":           patent_number,
        "priority_date":    priority_date,
        "charts":           charts,
        "reference_list":   list(all_refs.values()),
        "coverage_summary": coverage_summary,
        "gap_list":         all_gap_lims,
        "grace_annex":      all_grace_annex,
    }


# ─────────────────────────────────────────────────────────────
# Word document renderer
# ─────────────────────────────────────────────────────────────

def _render_word(drug_name: str, chart_data: dict, output_dir: Path) -> Optional[str]:
    """Render the chart data as an editable Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("[Step 9] python-docx not installed — skipping Word output")
        return None

    patent_number = chart_data["patent"]
    priority_date = chart_data["priority_date"]

    doc = Document()

    # Title
    title = doc.add_heading(f"Invalidity Claim Charts — {patent_number}", level=1)
    doc.add_paragraph(
        f"Drug: {drug_name}  |  Priority Date: {priority_date}  |  "
        f"Generated by IP Scope Engine"
    )

    # ── Charts ────────────────────────────────────────────────
    for chart in chart_data.get("charts", []):
        cn        = chart["claim_number"]
        ground_id = chart["ground_id"]
        basis     = chart["basis"]

        doc.add_heading(
            f"U.S. Patent No. {patent_number} — Claim {cn} ({ground_id}, §{basis})",
            level=2,
        )
        doc.add_paragraph(chart["basis_line"])
        doc.add_paragraph(f"Priority date: {priority_date}")
        if chart.get("combination_rationale"):
            doc.add_paragraph(f"Combination rationale: {chart['combination_rationale']}")

        # Two-column table
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Header row
        hdr = table.rows[0]
        hdr.cells[0].text = f"U.S. Patent No. {patent_number}, Claim {cn}"
        hdr.cells[1].text = ", ".join(chart["references"])
        for cell in hdr.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Limitation rows
        for row_data in chart["rows"]:
            row = table.add_row()
            # Left cell: limitation
            left_text = f"[{row_data['limitation_id']}] {row_data['limitation_text']}"
            row.cells[0].text = left_text

            # Right cell: evidence passages
            right_parts = []
            for cell_data in row_data.get("cells", []):
                ref_id = cell_data.get("reference_id", "")
                for p in cell_data.get("passages", []):
                    passage = p.get("passage_verbatim", "")
                    locus   = p.get("locus", "")
                    if locus:
                        right_parts.append(f'"{passage}" ({ref_id} at {locus})')
                    else:
                        right_parts.append(passage)
            row.cells[1].text = "\n\n".join(right_parts) if right_parts else "—"

        doc.add_paragraph()  # spacer

    # ── Reference List ────────────────────────────────────────
    doc.add_heading("Reference List", level=2)
    if chart_data.get("reference_list"):
        ref_table = doc.add_table(rows=1, cols=5)
        ref_table.style = "Table Grid"
        for i, h in enumerate(["Short Name", "Full Citation", "Pub Date", "Pre-Priority?", "Access"]):
            ref_table.rows[0].cells[i].text = h
            for run in ref_table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for ref in chart_data["reference_list"]:
            row = ref_table.add_row()
            row.cells[0].text = ref.get("short_name", "")
            row.cells[1].text = ref.get("full_citation", "")
            row.cells[2].text = ref.get("publication_date", "")
            row.cells[3].text = "Yes" if ref.get("pre_priority") else "No"
            row.cells[4].text = ref.get("access", "")

    # ── Coverage Summary ──────────────────────────────────────
    doc.add_heading("Coverage Summary", level=2)
    if chart_data.get("coverage_summary"):
        cov_table = doc.add_table(rows=1, cols=6)
        cov_table.style = "Table Grid"
        for i, h in enumerate(["Claim", "Ground", "References", "Covered/Total", "Coverage %", "Uncovered"]):
            cov_table.rows[0].cells[i].text = h
            for run in cov_table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        for cs in chart_data["coverage_summary"]:
            row = cov_table.add_row()
            row.cells[0].text = str(cs.get("claim", ""))
            row.cells[1].text = cs.get("ground", "")
            row.cells[2].text = ", ".join(cs.get("references", []))
            row.cells[3].text = cs.get("covered_total", "")
            row.cells[4].text = f"{cs.get('coverage_pct', 0)}%"
            row.cells[5].text = ", ".join(cs.get("uncovered_lim_ids", []))

    # ── Gap List ──────────────────────────────────────────────
    doc.add_heading("Gap List", level=2)
    if chart_data.get("gap_list"):
        for gap in chart_data["gap_list"]:
            flags = f" [{', '.join(gap['flags'])}]" if gap.get("flags") else ""
            rec   = f" → Recommended: {gap['recommended_source']}" if gap.get("recommended_source") else ""
            doc.add_paragraph(
                f"Claim {gap['claim_number']}, [{gap['limitation_id']}]: "
                f"{gap.get('limitation_text', '')}{flags}{rec}",
                style="List Bullet",
            )
            if gap.get("recommendation"):
                doc.add_paragraph(f"  {gap['recommendation']}")
        doc.add_paragraph("→ Direct paid / CAS / pharmacopoeia search for uncovered limitations.")
    else:
        doc.add_paragraph("No gaps — all limitations covered.")

    # ── Grace Annex ───────────────────────────────────────────
    doc.add_heading("Grace Annex", level=2)
    if chart_data.get("grace_annex"):
        for ga in chart_data["grace_annex"]:
            doc.add_paragraph(
                f"Claim {ga['claim_number']}, [{ga['limitation_id']}] ← {ga['reference_id']} "
                f"(within 12 months of priority — counsel admissibility review required)",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No grace-period-only limitations.")

    # Save
    safe_drug   = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name)
    safe_patent = re.sub(r"[^a-zA-Z0-9_-]", "_", patent_number)
    docx_path   = output_dir / f"{safe_drug}_{safe_patent}_claim_charts.docx"
    doc.save(str(docx_path))
    print(f"  → Word document   : {docx_path}")
    return str(docx_path)


# ─────────────────────────────────────────────────────────────
# Excel renderer
# ─────────────────────────────────────────────────────────────

def _render_excel(drug_name: str, chart_data: dict, output_dir: Path) -> Optional[str]:
    """
    Render the chart data as an editable Excel workbook.

    Sheets:
      1. Claim Charts     — one row per limitation per ground (the main chart)
      2. Coverage Summary — one row per claim per ground
      3. Reference List   — one row per cited reference
      4. Gap List         — uncovered limitations with recommended sources
      5. Grace Annex      — grace-period-only limitations for counsel review
    """
    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side, GradientFill
        )
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("[Step 9] openpyxl not installed — skipping Excel output")
        return None

    patent_number = chart_data["patent"]
    priority_date = chart_data["priority_date"]

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    # ── Shared styles ─────────────────────────────────────────
    FONT_NAME = "Arial"

    def _hdr_font(bold=True):
        return Font(name=FONT_NAME, bold=bold, size=10, color="FFFFFF")

    def _body_font(bold=False):
        return Font(name=FONT_NAME, bold=bold, size=10)

    def _fill(hex_color):
        return PatternFill("solid", fgColor=hex_color)

    def _border():
        thin = Side(style="thin", color="BFBFBF")
        return Border(left=thin, right=thin, top=thin, bottom=thin)

    def _wrap():
        return Alignment(wrap_text=True, vertical="top")

    HEADER_FILL_DARK  = _fill("1F4E79")   # dark blue   — sheet / section headers
    HEADER_FILL_MID   = _fill("2E75B6")   # mid blue    — column headers
    HEADER_FILL_102   = _fill("375623")   # dark green  — §102 ground rows
    HEADER_FILL_103   = _fill("7F6000")   # dark amber  — §103 ground rows
    FILL_COVERED      = _fill("E2EFDA")   # light green — covered limitation
    FILL_GAP          = _fill("FCE4D6")   # light red   — not disclosed / gap
    FILL_GRACE        = _fill("FFF2CC")   # light amber — grace period

    def _write_col_headers(ws, headers: list[str], row: int, fill=None) -> None:
        for col, h in enumerate(headers, 1):
            c = ws.cell(row=row, column=col, value=h)
            c.font      = _hdr_font()
            c.fill      = fill or HEADER_FILL_MID
            c.border    = _border()
            c.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")

    def _body_cell(ws, row, col, value, fill=None, bold=False, wrap=True):
        c = ws.cell(row=row, column=col, value=str(value) if value is not None else "")
        c.font      = _body_font(bold=bold)
        c.border    = _border()
        c.alignment = _wrap() if wrap else Alignment(vertical="top")
        if fill:
            c.fill = fill
        return c

    def _set_col_widths(ws, widths: list[int]) -> None:
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

    def _title_row(ws, text: str, row: int, n_cols: int) -> None:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=n_cols)
        c = ws.cell(row=row, column=1, value=text)
        c.font      = Font(name=FONT_NAME, bold=True, size=11, color="FFFFFF")
        c.fill      = HEADER_FILL_DARK
        c.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[row].height = 22

    # ══════════════════════════════════════════════════════════
    # Sheet 1: Claim Charts
    # ══════════════════════════════════════════════════════════
    ws1 = wb.create_sheet("Claim Charts")
    _set_col_widths(ws1, [12, 8, 12, 8, 55, 18, 70, 40, 40])

    CHART_HEADERS = [
        "Patent No.", "Claim", "Ground", "Basis",
        "Limitation (verbatim)", "Limitation ID",
        "Prior Art Passage(s)", "Reference(s)", "Locus / Citation",
    ]

    r = 1
    _title_row(ws1,
        f"Invalidity Claim Charts  |  {patent_number}  |  Drug: {drug_name}  |  Priority: {priority_date}",
        r, len(CHART_HEADERS))
    r += 1
    _write_col_headers(ws1, CHART_HEADERS, r)
    r += 1

    for chart in chart_data.get("charts", []):
        cn        = chart["claim_number"]
        ground_id = chart["ground_id"]
        basis     = chart["basis"]
        refs      = ", ".join(chart.get("references", []))
        g_fill    = HEADER_FILL_102 if basis == "102" else HEADER_FILL_103

        # Ground header row (merged basis line)
        ws1.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(CHART_HEADERS))
        gc = ws1.cell(row=r, column=1, value=chart.get("basis_line", ""))
        gc.font      = Font(name=FONT_NAME, bold=True, size=10, color="FFFFFF")
        gc.fill      = g_fill
        gc.alignment = Alignment(wrap_text=True, vertical="center")
        ws1.row_dimensions[r].height = 18
        r += 1

        for row_data in chart.get("rows", []):
            lid    = row_data.get("limitation_id", "")
            ltext  = row_data.get("limitation_text", "")
            cells  = row_data.get("cells", [])

            # Determine if this limitation is covered or a gap
            is_gap = any(
                "Not disclosed" in (p.get("passage_verbatim", ""))
                for cd in cells for p in cd.get("passages", [])
            )
            row_fill = FILL_GAP if is_gap else FILL_COVERED

            # Build passage / reference / locus strings
            passages_parts = []
            ref_parts      = []
            locus_parts    = []
            for cd in cells:
                ref_id = cd.get("reference_id", "")
                for p in cd.get("passages", []):
                    passage = p.get("passage_verbatim", "")
                    locus   = p.get("locus", "")
                    if passage and "Not disclosed" not in passage:
                        label = f"[{ref_id}] " if len(cells) > 1 else ""
                        passages_parts.append(f'{label}"{passage}"')
                        ref_parts.append(ref_id)
                        locus_parts.append(f"{ref_id} at {locus}" if locus else ref_id)

            passage_text = "\n\n".join(passages_parts) if passages_parts else "Not disclosed — see Gap List"
            ref_text     = "; ".join(dict.fromkeys(ref_parts)) or "—"
            locus_text   = "\n".join(locus_parts) or "—"

            _body_cell(ws1, r, 1, patent_number)
            _body_cell(ws1, r, 2, cn)
            _body_cell(ws1, r, 3, ground_id)
            _body_cell(ws1, r, 4, f"§{basis}")
            _body_cell(ws1, r, 5, ltext, fill=row_fill)
            _body_cell(ws1, r, 6, lid)
            _body_cell(ws1, r, 7, passage_text, fill=row_fill)
            _body_cell(ws1, r, 8, ref_text)
            _body_cell(ws1, r, 9, locus_text)
            ws1.row_dimensions[r].height = max(40, min(120, 15 * (passage_text.count("\n") + 2)))
            r += 1

        r += 1  # blank row between grounds

    ws1.freeze_panes = "A3"

    # ══════════════════════════════════════════════════════════
    # Sheet 2: Coverage Summary
    # ══════════════════════════════════════════════════════════
    ws2 = wb.create_sheet("Coverage Summary")
    _set_col_widths(ws2, [12, 8, 8, 40, 15, 12, 12, 50, 50])

    COV_HEADERS = [
        "Patent No.", "Claim", "Ground", "Reference(s)",
        "Covered / Total", "Coverage %", "Basis",
        "Uncovered Limitation IDs", "Combination Rationale",
    ]
    r = 1
    _title_row(ws2, f"Coverage Summary  |  {patent_number}  |  Priority: {priority_date}",
               r, len(COV_HEADERS))
    r += 1
    _write_col_headers(ws2, COV_HEADERS, r)
    r += 1

    for cs in chart_data.get("coverage_summary", []):
        basis  = cs.get("ground", "").replace("§", "")
        g_fill = FILL_COVERED if cs.get("coverage_pct", 0) == 100.0 else None
        _body_cell(ws2, r, 1, patent_number)
        _body_cell(ws2, r, 2, cs.get("claim", ""))
        _body_cell(ws2, r, 3, cs.get("ground", ""))
        _body_cell(ws2, r, 4, ", ".join(cs.get("references", [])))
        _body_cell(ws2, r, 5, cs.get("covered_total", ""))
        _body_cell(ws2, r, 6, f"{cs.get('coverage_pct', 0)}%", fill=g_fill)
        _body_cell(ws2, r, 7, cs.get("ground", ""))
        _body_cell(ws2, r, 8, ", ".join(cs.get("uncovered_lim_ids", [])) or "—")
        _body_cell(ws2, r, 9, cs.get("combination_rationale", "—"))
        ws2.row_dimensions[r].height = 30
        r += 1

    ws2.freeze_panes = "A3"

    # ══════════════════════════════════════════════════════════
    # Sheet 3: Reference List
    # ══════════════════════════════════════════════════════════
    ws3 = wb.create_sheet("Reference List")
    _set_col_widths(ws3, [25, 60, 15, 14, 25])

    REF_HEADERS = [
        "Short Name", "Full Citation", "Publication Date",
        "Pre-Priority?", "Access",
    ]
    r = 1
    _title_row(ws3, f"Reference List  |  {patent_number}", r, len(REF_HEADERS))
    r += 1
    _write_col_headers(ws3, REF_HEADERS, r)
    r += 1

    for ref in chart_data.get("reference_list", []):
        pre = "Yes" if ref.get("pre_priority") else "No"
        pre_fill = FILL_COVERED if ref.get("pre_priority") else FILL_GAP
        _body_cell(ws3, r, 1, ref.get("short_name", ""), bold=True)
        _body_cell(ws3, r, 2, ref.get("full_citation", ""))
        _body_cell(ws3, r, 3, ref.get("publication_date", ""))
        _body_cell(ws3, r, 4, pre, fill=pre_fill)
        _body_cell(ws3, r, 5, ref.get("access", "publicly accessible"))
        ws3.row_dimensions[r].height = 25
        r += 1

    ws3.freeze_panes = "A3"

    # ══════════════════════════════════════════════════════════
    # Sheet 4: Gap List
    # ══════════════════════════════════════════════════════════
    ws4 = wb.create_sheet("Gap List")
    _set_col_widths(ws4, [8, 12, 60, 30, 30, 50])

    GAP_HEADERS = [
        "Claim", "Limitation ID", "Limitation (verbatim)",
        "Flags", "Recommended Source", "Recommendation",
    ]
    r = 1
    _title_row(ws4,
        f"Gap List — Uncovered Limitations  |  {patent_number}  |  "
        f"→ Direct paid / CAS / pharmacopoeia search",
        r, len(GAP_HEADERS))
    r += 1
    _write_col_headers(ws4, GAP_HEADERS, r)
    r += 1

    for gap in chart_data.get("gap_list", []):
        _body_cell(ws4, r, 1, gap.get("claim_number", ""))
        _body_cell(ws4, r, 2, gap.get("limitation_id", ""), bold=True, fill=FILL_GAP)
        _body_cell(ws4, r, 3, gap.get("limitation_text", ""))
        _body_cell(ws4, r, 4, "; ".join(gap.get("flags", [])) or "—")
        _body_cell(ws4, r, 5, gap.get("recommended_source", "—"), fill=FILL_GRACE)
        _body_cell(ws4, r, 6, gap.get("recommendation", "—"))
        ws4.row_dimensions[r].height = 35
        r += 1

    if r == 3:  # no gaps
        ws4.cell(row=r, column=1, value="No gaps — all limitations covered.").font = _body_font()

    ws4.freeze_panes = "A3"

    # ══════════════════════════════════════════════════════════
    # Sheet 5: Grace Annex
    # ══════════════════════════════════════════════════════════
    ws5 = wb.create_sheet("Grace Annex")
    _set_col_widths(ws5, [8, 12, 35, 50])

    GRACE_HEADERS = [
        "Claim", "Limitation ID", "Reference (Grace Period)",
        "Note",
    ]
    r = 1
    _title_row(ws5,
        f"Grace Annex — Counsel Admissibility Review Required  |  {patent_number}",
        r, len(GRACE_HEADERS))
    r += 1
    _write_col_headers(ws5, GRACE_HEADERS, r)
    r += 1

    for ga in chart_data.get("grace_annex", []):
        _body_cell(ws5, r, 1, ga.get("claim_number", ""))
        _body_cell(ws5, r, 2, ga.get("limitation_id", ""), bold=True, fill=FILL_GRACE)
        _body_cell(ws5, r, 3, ga.get("reference_id", ""))
        _body_cell(ws5, r, 4,
            "Published within 12 months of priority date — "
            "admissibility case-by-case; counsel to confirm")
        ws5.row_dimensions[r].height = 30
        r += 1

    if r == 3:
        ws5.cell(row=r, column=1, value="No grace-period-only limitations.").font = _body_font()

    ws5.freeze_panes = "A3"

    # ── Save ──────────────────────────────────────────────────
    safe_drug   = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name)
    safe_patent = re.sub(r"[^a-zA-Z0-9_-]", "_", patent_number)
    xlsx_path   = output_dir / f"{safe_drug}_{safe_patent}_claim_charts.xlsx"
    wb.save(str(xlsx_path))
    print(f"  → Excel workbook  : {xlsx_path}")
    return str(xlsx_path)

def process_patent(
    drug_name:      str,
    patent_number:  str,
    output_dir:     Path,
) -> Optional[dict]:
    """Assemble final charts for one patent."""
    inputs = _load_all_inputs(drug_name, patent_number)

    skeleton = inputs["skeleton"]
    evidence = inputs["evidence"]
    grounds  = inputs["grounds"]

    if not skeleton:
        print(f"[Step 9] ⚠ No step7 skeleton for {patent_number}")
        return None
    if not evidence:
        print(f"[Step 9] ⚠ No step8a evidence for {patent_number}")
        return None
    if not grounds:
        print(f"[Step 9] ⚠ No step8b grounds for {patent_number}")
        return None

    priority_date = skeleton.get("priority_date", evidence.get("priority_date", "Unknown"))

    print(f"\n[Step 9] Assembling charts for {patent_number} (priority: {priority_date})...")

    chart_data = _assemble_charts(
        patent_number = patent_number,
        priority_date = priority_date,
        skeleton      = skeleton,
        evidence      = evidence,
        grounds_data  = grounds,
    )

    # Write JSON
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_drug   = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name)
    safe_patent = re.sub(r"[^a-zA-Z0-9_-]", "_", patent_number)

    json_path = output_dir / f"{safe_drug}_{safe_patent}_final_charts.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(chart_data, f, indent=2)
    print(f"  → Charts JSON     : {json_path}")

    # Write Word
    _render_word(drug_name, chart_data, output_dir)

    # Write Excel
    _render_excel(drug_name, chart_data, output_dir)

    # Summary
    n_charts = len(chart_data.get("charts", []))
    n_refs   = len(chart_data.get("reference_list", []))
    n_gaps   = len(chart_data.get("gap_list", []))
    n_grace  = len(chart_data.get("grace_annex", []))
    print(f"  Summary: {n_charts} chart(s), {n_refs} reference(s), "
          f"{n_gaps} gap(s), {n_grace} grace item(s)")

    return chart_data


# ─────────────────────────────────────────────────────────────
# Main runner
# ─────────────────────────────────────────────────────────────

def run_for_drug(
    drug_name:     str,
    patent_filter: Optional[str] = None,
    rerun_all:     bool          = False,
    output_dir:    Path          = STEP9_OUTPUT_DIR,
) -> list[dict]:
    _ensure_step8b(drug_name, force=rerun_all)

    # Discover all patents from step8b grounds files
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", drug_name)
    grounds_files = sorted(STEP8B_OUTPUT_DIR.glob(f"{safe}_*_grounds.json"))
    grounds_files = [f for f in grounds_files if "_all_grounds" not in f.name]

    patent_numbers = []
    for f in grounds_files:
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            pn   = data.get("patent_number", "")
            if patent_filter and pn.upper() != patent_filter.upper():
                continue
            patent_numbers.append(pn)
        except Exception:
            continue

    if not patent_numbers:
        print(f"[Step 9] No patents found for '{drug_name}'.")
        return []

    print(f"[Step 9] Assembling charts for {len(patent_numbers)} patent(s)...")

    results = []
    for pn in patent_numbers:
        result = process_patent(drug_name, pn, output_dir)
        if result:
            results.append(result)

    if results:
        combined_path = output_dir / f"{safe}_all_final_charts.json"
        with open(combined_path, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=2)
        print(f"\n  → Combined JSON   : {combined_path}")

    return results


# ─────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Step 9 – Invalidity Claim Chart Assembler (final deliverable)"
    )
    parser.add_argument("--drug",       "-d", required=True)
    parser.add_argument("--patent",     "-p", default=None)
    parser.add_argument("--rerun_all",  action="store_true",
                        help="Cascade re-run steps 7 → 8a → 8b → 9")
    parser.add_argument("--output_dir", default=str(STEP9_OUTPUT_DIR))
    args = parser.parse_args()

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"[Step 9] Drug       : {args.drug}")
    print(f"[Step 9] Patent     : {args.patent or 'all'}")
    print(f"[Step 9] Output     : {output_dir.resolve()}")

    results = run_for_drug(
        drug_name     = args.drug,
        patent_filter = args.patent,
        rerun_all     = args.rerun_all,
        output_dir    = output_dir,
    )

    total_charts = sum(len(r.get("charts", [])) for r in results)
    total_gaps   = sum(len(r.get("gap_list", [])) for r in results)
    print(f"\n[Step 9] Done. {total_charts} chart(s) rendered, {total_gaps} gap(s) remaining.")

    if not results:
        sys.exit(1)


if __name__ == "__main__":
    main()
